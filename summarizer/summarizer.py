from tqdm.auto import tqdm
import csv
import textwrap
from pathlib import Path
from docx import Document
from bs4 import BeautifulSoup
from markdown import markdown


def chunk_text(text: str, max_tokens: int = 2000) -> list:
    """Split text into chunks of approximately max_tokens"""
    # Simple splitting by newlines first
    paragraphs = text.split('\n')
    chunks = []
    current_chunk = []
    current_length = 0
    
    for para in paragraphs:
        # Rough estimation of tokens (words + some overhead)
        para_length = len(para.split())
        if current_length + para_length > max_tokens and current_chunk:
            chunks.append('\n'.join(current_chunk))
            current_chunk = [para]
            current_length = para_length
        else:
            current_chunk.append(para)
            current_length += para_length
    
    if current_chunk:
        chunks.append('\n'.join(current_chunk))
    return chunks


def get_chunk_summary(chunk: str) -> str:
    """Get summary for a single chunk"""
    prompt = f"""Please summarize the following Thai conversation segment into key points:

- Extract main discussion points with some key details
- Note any decisions or action items
- Highlight suggestions or feedback

Text segment:
{chunk}

Please provide the summary in Thai language."""
    return prompt


def combine_summaries(summaries: list) -> str:
    """Combine multiple chunk summaries into final summary"""
    combined_text = "\n\n".join(summaries)
    
    final_prompt = f"""Generate a detailed Thai report summarizing the contents of an ASR. The report should be well-structured and organized by a chronological timeline with report structure
    1. Introduction (purpose of the conversation, key participants if mentioned, timeline of the conversation)
    2. Main discussion points (Divide the report into clear subtopics reflecting the flow of discussion,  key points with elaborated details to clearify the points)
    3. Follow-up Questions and Action Items (outstanding questions or required actions in bullet points, critical follow-up points; no reccommendations)
    4. Conclusions and reccommendations (main takeaways from the conversation, suggested next steps or recommendations based on the discussion)
    Segment summaries 
    {combined_text}

    Please provide the final summary in Thai language.
    Formatting Guidelines:
    - Use concise yet informative language.
    - Ensure clarity by breaking down complex information into simple terms.
    - Maintain a formal and professional tone throughout the report."""
    
    return final_prompt


def markdown_to_docx(markdown_text: str, output_filename: str = 'output.docx') -> None:
   """
   Convert markdown text to docx file
   
   Args:
       markdown_text (str): Input markdown text
       output_filename (str): Output docx filename
   """
   # Convert markdown to HTML
   html = markdown(markdown_text)
   
   # Create document object
   doc = Document()
   
   # Parse HTML
   soup = BeautifulSoup(html, 'html.parser')
   
   # Process each element
   for element in soup.descendants:
       if element.name == 'h1':
           doc.add_heading(element.text, 0)
       elif element.name == 'h2':
           doc.add_heading(element.text, 1)
       elif element.name == 'h3':
           doc.add_heading(element.text, 2)
       elif element.name == 'p':
           doc.add_paragraph(element.text)
       elif element.name == 'li':
           doc.add_paragraph(element.text, style='List Bullet')
       elif element.name == 'code':
           p = doc.add_paragraph()
           p.add_run(element.text).font.name = 'Courier'
           
   # Save document
   doc.save(output_filename)


def summarise_from_file(model, file: str, add_transcript: bool = True, output_path: str = "summary_report.docx"):
    file = Path(file)

    # read CSV file
    if file.name.endswith('.csv'):
        with open(file.name, 'r') as f:
            csv_reader = csv.reader(f)
            next(csv_reader)  # Skip header
            transcript = "\n".join(row[1] for row in csv_reader)
            chunks = chunk_text(transcript)
    # read general text file
    else:
        with file.open('r') as f:
            text = f.read().strip()
        num_rows = 10
        chunk_size = max(1, len(text) // num_rows)
        chunks = textwrap.wrap(text, chunk_size)

    chunk_summaries = []
    for chunk in tqdm(chunks):
        summary_prompt = get_chunk_summary(chunk)  # prompt preparation
        summary = model.generate_content(summary_prompt).text
        chunk_summaries.append(summary)
    final_summary_prompt = combine_summaries(chunk_summaries)
    final_summary = model.generate_content(final_summary_prompt).text

    # Add transcript to final summary
    if add_transcript:
        final_summary += f"\n\n\n## ผลการถอดความ:\n{transcript}"

    # Save to Word file
    markdown_to_docx(final_summary, output_path)

    return final_summary, output_path
